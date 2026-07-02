"""Embed diagram images into Word documents (PNG for display compatibility)."""
from __future__ import annotations

import logging
from pathlib import Path

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

log = logging.getLogger(__name__)


def add_centered_png_figure(
    doc: Document,
    png_path: str,
    caption: str,
    width_inches: float = 6.0,
) -> None:
    """Insert a PNG diagram — universally visible in Microsoft Word."""
    if not Path(png_path).is_file():
        log.warning("PNG figure missing, skipping embed: %s", png_path)
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(png_path, width=Inches(width_inches))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "Caption" in doc.styles:
        cap.style = "Caption"


def add_centered_svg_figure(
    doc: Document,
    svg_path: str,
    caption: str,
    width_inches: float = 6.0,
    png_path: str | None = None,
) -> None:
    """
    Insert a diagram figure into Word.

    Word does not reliably render embedded SVG (image/svg+xml) — diagrams appear
    blank even when the parts exist in the package. Prefer PNG when available.
    """
    png_candidate = png_path or str(Path(svg_path).with_suffix(".png"))
    if Path(png_candidate).is_file():
        add_centered_png_figure(doc, png_candidate, caption, width_inches)
        return

    log.warning(
        "No PNG available for %s — SVG-only embed is not visible in most Word versions. "
        "Regenerate with PNG pipeline.",
        svg_path,
    )
