"""
Embed diagrams as native Word DrawingML shapes (editable rectangles, text, lines).

These are NOT images (PNG/SVG). Each box and connector is a separate Word shape
you can click, move, recolor, and edit in Microsoft Word.
"""
from __future__ import annotations

import logging
import math
from typing import Dict, List, Tuple

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

log = logging.getLogger(__name__)

EMU_PER_INCH = 914400
REF_W = 26.0
REF_H = 15.0
_SHAPE_ID = 100

# Full namespace block for Word DrawingML shape groups (wps/wpg not in docx nsmap)
_DRAWML_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
)


def _next_shape_id() -> int:
    global _SHAPE_ID
    _SHAPE_ID += 1
    return _SHAPE_ID


def reset_shape_ids() -> None:
    global _SHAPE_ID
    _SHAPE_ID = 100


def _hex6(color: str) -> str:
    c = (color or "#000000").lstrip("#")
    return c[:6] if len(c) >= 6 else "000000"


def _esc(text: str) -> str:
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _inch_to_group_emu(
    x_in: float, y_in: float, w_in: float, h_in: float, group_cx: int, group_cy: int
) -> Tuple[int, int, int, int]:
    """Map layout inch coords (center x,y) to group EMU top-left + size."""
    x0 = int((x_in - w_in / 2) / REF_W * group_cx)
    y0 = int((y_in - h_in / 2) / REF_H * group_cy)
    w = max(int(w_in / REF_W * group_cx), 9144)
    h = max(int(h_in / REF_H * group_cy), 9144)
    return x0, y0, w, h


def _textbox_paragraphs(text: str, font_size: int = 9, bold: bool = False) -> str:
    lines = str(text or " ").split("\n")[:6]
    parts: List[str] = []
    for line in lines:
        weight = '<w:b/>' if bold else ''
        parts.append(
            f'<w:p><w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
            f'<w:r><w:rPr><w:sz w:val="{font_size * 2}"/>{weight}</w:rPr>'
            f'<w:t xml:space="preserve">{_esc(line)}</w:t></w:r></w:p>'
        )
    return "".join(parts)


def _rect_shape(
    x0: int, y0: int, w: int, h: int,
    fill: str, border: str, text: str,
    text_color: str = "000000",
    shape_id: int = 1,
    round_rect: bool = True,
) -> str:
    geom = "roundRect" if round_rect else "rect"
    txbx = _textbox_paragraphs(text) if text.strip() else '<w:p/>'
    return f"""
    <wps:wsp>
      <wps:cNvSpPr><a:spLocks noChangeAspect="1"/></wps:cNvSpPr>
      <wps:spPr>
        <a:xfrm>
          <a:off x="{x0}" y="{y0}"/>
          <a:ext cx="{w}" cy="{h}"/>
        </a:xfrm>
        <a:prstGeom prst="{geom}"><a:avLst/></a:prstGeom>
        <a:solidFill><a:srgbClr val="{_hex6(fill)}"/></a:solidFill>
        <a:ln w="12700"><a:solidFill><a:srgbClr val="{_hex6(border)}"/></a:solidFill></a:ln>
      </wps:spPr>
      <wps:txbx><w:txbxContent>{txbx}</w:txbxContent></wps:txbx>
      <wps:bodyPr wrap="square" lIns="36000" tIns="18000" rIns="36000" bIns="18000" anchor="ctr"/>
    </wps:wsp>"""


def _line_shape(x1: int, y1: int, x2: int, y2: int, color: str, shape_id: int) -> str:
    length = int(math.hypot(x2 - x1, y2 - y1))
    if length < 5000:
        return ""
    rot = int(math.degrees(math.atan2(y2 - y1, x2 - x1)) * 60000)
    return f"""
    <wps:wsp>
      <wps:cNvSpPr/>
      <wps:spPr>
        <a:xfrm rot="{rot}">
          <a:off x="{x1}" y="{y1}"/>
          <a:ext cx="{length}" cy="0"/>
        </a:xfrm>
        <a:prstGeom prst="line"><a:avLst/></a:prstGeom>
        <a:ln w="12700"><a:solidFill><a:srgbClr val="{_hex6(color)}"/></a:solidFill></a:ln>
      </wps:spPr>
      <wps:bodyPr/>
    </wps:wsp>"""


def layout_to_drawingml_group(layout: dict, width_inches: float = 6.0):
    """Build an inline grouped DrawingML element from a layout dict."""
    group_cx = int(width_inches * EMU_PER_INCH)
    group_cy = int(group_cx * REF_H / REF_W)
    doc_pr_id = _next_shape_id()

    child_shapes: List[str] = []
    centers: Dict[str, Tuple[int, int]] = {}
    bounds: Dict[str, Tuple[int, int, int, int]] = {}

    for node in layout.get("nodes", []):
        x0, y0, w, h = _inch_to_group_emu(
            node.get("x", 0), node.get("y", 0),
            node.get("w", 1), node.get("h", 0.5),
            group_cx, group_cy,
        )
        centers[node["id"]] = (x0 + w // 2, y0 + h // 2)
        bounds[node["id"]] = (x0, y0, x0 + w, y0 + h)
        sid = _next_shape_id()
        child_shapes.append(_rect_shape(
            x0, y0, w, h,
            fill=node.get("fill", "#FFFFFF"),
            border=node.get("border", "#666666"),
            text=str(node.get("text", "")),
            text_color=node.get("text_color", "#000000"),
            shape_id=sid,
        ))

    for edge in layout.get("edges", []):
        src = bounds.get(edge.get("from", ""))
        dst = bounds.get(edge.get("to", ""))
        if not src or not dst:
            continue
        x1, y1 = src[2], (src[1] + src[3]) // 2
        x2, y2 = dst[0], (dst[1] + dst[3]) // 2
        line = _line_shape(x1, y1, x2, y2, edge.get("color", "#666666"), _next_shape_id())
        if line:
            child_shapes.append(line)

    children_xml = "".join(child_shapes)

    xml = f"""
    <w:drawing {_DRAWML_NS}>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{group_cx}" cy="{group_cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="{doc_pr_id}" name="Diagram_{doc_pr_id}" descr="{_esc(layout.get('title', 'Diagram'))}"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
            <wpg:wgp>
              <wpg:cNvGrpSpPr><a:grpSpLocks noChangeAspect="1"/></wpg:cNvGrpSpPr>
              <wpg:grpSpPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{group_cx}" cy="{group_cy}"/>
                  <a:chOff x="0" y="0"/>
                  <a:chExt cx="{group_cx}" cy="{group_cy}"/>
                </a:xfrm>
              </wpg:grpSpPr>
              {children_xml}
            </wpg:wgp>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>"""
    return parse_xml(xml)


def add_centered_drawingml_figure(
    doc: Document,
    layout: dict,
    caption: str,
    width_inches: float = 6.0,
) -> None:
    """Insert an editable native Word shape diagram (DrawingML group)."""
    if not layout or not layout.get("nodes"):
        log.warning("Empty layout — skipping DrawingML figure for %s", caption)
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run._r.append(layout_to_drawingml_group(layout, width_inches))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "Caption" in doc.styles:
        cap.style = "Caption"

    log.debug("Embedded DrawingML figure: %s (%d shapes)", caption, len(layout.get("nodes", [])))
