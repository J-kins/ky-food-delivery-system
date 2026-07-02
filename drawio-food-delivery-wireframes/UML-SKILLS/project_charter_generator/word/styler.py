"""Corporate Word styling helpers."""
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


DARK_BLUE = RGBColor(0x1A, 0x23, 0x7E)
PRIMARY_BLUE = RGBColor(0x15, 0x65, 0xC0)


def apply_corporate_styles(doc) -> None:
    """Apply base styles when no template is available."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 16), (2, 14), (3, 12)):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.name = "Calibri"
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = DARK_BLUE if level <= 2 else PRIMARY_BLUE


def add_centered_figure(doc, image_path: str, caption: str, width_inches: float = 6.0) -> None:
    from docx.shared import Inches

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "Caption" in doc.styles:
        cap.style = "Caption"
