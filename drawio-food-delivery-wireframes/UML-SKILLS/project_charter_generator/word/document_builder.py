"""Build the full 13-section Project Charter Word document."""
from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from word.drawingml_inserter import add_centered_drawingml_figure, reset_shape_ids
from word.sections import add_bullet_list, add_table
from word.styler import apply_corporate_styles
from word.svg_inserter import add_centered_svg_figure


def _figure_width(payload: Dict[str, Any]) -> float:
    wd = payload.get("word_document") or {}
    return float(wd.get("figure_width_inches", 6.0))


def _open_document() -> Document:
    template_candidates = [
        Path(__file__).resolve().parent.parent / "templates" / "word_template.docx",
        Path("templates/word_template.docx"),
    ]
    for path in template_candidates:
        if path.is_file():
            try:
                return Document(str(path))
            except Exception:
                continue
    return Document()


def _add_figure(doc, diagrams: Dict[str, Any], key: str, caption: str, width: float) -> None:
    artifact = diagrams.get(key)
    if not artifact:
        return
    cap = artifact.get("caption", caption) if isinstance(artifact, dict) else caption
    layout = artifact.get("layout") if isinstance(artifact, dict) else None

    # Prefer native editable DrawingML shapes (click-to-edit boxes in Word)
    if layout and layout.get("nodes"):
        add_centered_drawingml_figure(doc, layout, cap, width)
        return

    # Fallback: PNG raster if layout unavailable
    svg_path = artifact.get("svg") if isinstance(artifact, dict) else artifact
    png_path = artifact.get("png") if isinstance(artifact, dict) else None
    if svg_path:
        png_candidate = png_path or str(Path(str(svg_path)).with_suffix(".png"))
        if Path(png_candidate).is_file() or Path(str(svg_path)).is_file():
            add_centered_svg_figure(
                doc, str(svg_path), cap, width,
                png_path=png_candidate if Path(png_candidate).is_file() else None,
            )


def build_word_document(payload: Dict[str, Any], diagrams: Dict[str, Any], output_path: str) -> str:
    reset_shape_ids()
    doc = _open_document()
    apply_corporate_styles(doc)

    project = payload.get("project", {})
    vision = payload.get("vision", {})
    scope = payload.get("scope", {})
    budget = payload.get("budget") or {}
    fig_w = _figure_width(payload)

    # Title page
    title = doc.add_paragraph("PROJECT CHARTER", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(project.get("name", ""), style="Title")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        f"Version {project.get('version', '1.0')} | "
        f"{project.get('start_date', '')} – {project.get('end_date', '')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading("1. EXECUTIVE SUMMARY", level=1)
    summary = (
        f"{project.get('name', 'This project')} is sponsored by {project.get('sponsor', 'TBD')} "
        f"and managed by {project.get('manager', 'TBD')}. "
        f"{vision.get('statement', '')}"
    )
    doc.add_paragraph(summary.strip())

    # 2. Project Overview
    doc.add_heading("2. PROJECT OVERVIEW", level=1)
    overview_rows = [
        ("Project Name", project.get("name", "")),
        ("Project Sponsor", project.get("sponsor", "")),
        ("Project Manager", project.get("manager", "")),
        ("Department/Division", project.get("department", "")),
        ("Start Date", project.get("start_date", "")),
        ("End Date", project.get("end_date", "")),
    ]
    add_table(doc, ["Field", "Value"], overview_rows)

    # 3. Vision & Objectives
    doc.add_heading("3. VISION & OBJECTIVES", level=1)
    doc.add_heading("3.1 Vision Statement", level=2)
    p = doc.add_paragraph(vision.get("statement", ""))
    p.runs[0].bold = True if p.runs else None
    doc.add_heading("3.2 Mission Statement", level=2)
    doc.add_paragraph(vision.get("mission", ""))
    doc.add_heading("3.3 SMART Objectives", level=2)
    obj_rows = [
        (o.get("id", ""), o.get("description", ""), o.get("measurable_criteria", ""))
        for o in payload.get("objectives", [])
    ]
    add_table(doc, ["ID", "Objective", "Measurable Criteria"], obj_rows)

    # 4. Scope
    doc.add_heading("4. SCOPE", level=1)
    doc.add_heading("4.1 In-Scope", level=2)
    add_bullet_list(doc, scope.get("in_scope", []))
    doc.add_heading("4.2 Out-of-Scope", level=2)
    add_bullet_list(doc, scope.get("out_of_scope", []))
    if scope.get("boundaries"):
        doc.add_heading("4.3 Boundaries", level=2)
        doc.add_paragraph(scope["boundaries"])
    doc.add_heading("4.4 Scope Diagram", level=2)
    _add_figure(doc, diagrams, "scope_boundary", "Figure 1: Scope Boundaries", fig_w)

    # 5. Stakeholders
    doc.add_heading("5. STAKEHOLDERS", level=1)
    doc.add_heading("5.1 Stakeholder Register", level=2)
    sh_rows = [
        (
            s.get("id", ""),
            s.get("name", ""),
            s.get("role", ""),
            s.get("organization", ""),
            s.get("power", ""),
            s.get("interest", ""),
        )
        for s in payload.get("stakeholders", [])
    ]
    add_table(doc, ["ID", "Name", "Role", "Organization", "Power", "Interest"], sh_rows)
    doc.add_heading("5.2 Stakeholder Matrix", level=2)
    _add_figure(doc, diagrams, "stakeholder_matrix", "Figure 2: Power-Interest Matrix", fig_w)

    # 6. System Context
    doc.add_heading("6. SYSTEM CONTEXT", level=1)
    ctx = payload.get("diagrams", {}).get("system_context") or {}
    doc.add_paragraph(ctx.get("description", "System context diagram (when provided)."))
    doc.add_heading("6.2 Context Diagram", level=2)
    _add_figure(doc, diagrams, "system_context", "Figure: System Context", fig_w)

    # 7. Project Organization
    doc.add_heading("7. PROJECT ORGANIZATION", level=1)
    doc.add_heading("7.1 Team Structure", level=2)
    team_rows = [
        (m.get("id", ""), m.get("name", ""), m.get("role", ""), m.get("reports_to") or "—")
        for m in payload.get("team", [])
    ]
    add_table(doc, ["ID", "Name", "Role", "Reports To"], team_rows)
    doc.add_heading("7.2 Org Chart", level=2)
    _add_figure(doc, diagrams, "org_chart", "Figure 3: Project Organization", fig_w)

    # 8. Constraints & Assumptions
    doc.add_heading("8. CONSTRAINTS & ASSUMPTIONS", level=1)
    doc.add_heading("8.1 Constraints", level=2)
    add_bullet_list(doc, payload.get("constraints", []))
    doc.add_heading("8.2 Assumptions", level=2)
    add_bullet_list(doc, payload.get("assumptions", []))

    # 9. Risks
    doc.add_heading("9. RISKS", level=1)
    doc.add_heading("9.1 Risk Register", level=2)
    risk_rows = [
        (
            r.get("id", ""),
            r.get("description", ""),
            r.get("likelihood", ""),
            r.get("impact", ""),
            r.get("mitigation", ""),
        )
        for r in payload.get("risks", [])
    ]
    add_table(doc, ["ID", "Risk", "Likelihood", "Impact", "Mitigation"], risk_rows)
    doc.add_heading("9.2 Problem Tree Analysis", level=2)
    _add_figure(doc, diagrams, "problem_tree", "Figure 4: Problem Tree", fig_w)
    doc.add_heading("9.3 Risk Matrix", level=2)
    _add_figure(doc, diagrams, "risk_matrix", "Figure 5: Risk Matrix", fig_w)

    # 10. Milestones
    doc.add_heading("10. MILESTONES", level=1)
    doc.add_heading("10.1 Milestone Schedule", level=2)
    ms_rows = [
        (m.get("id", ""), m.get("name", ""), m.get("date", ""), m.get("deliverable", ""))
        for m in payload.get("milestones", [])
    ]
    add_table(doc, ["ID", "Milestone", "Date", "Deliverable"], ms_rows)
    doc.add_heading("10.2 Timeline", level=2)
    _add_figure(doc, diagrams, "milestone_timeline", "Figure 6: Milestone Timeline", fig_w)

    # 11. Budget
    doc.add_heading("11. BUDGET", level=1)
    if budget:
        bd = (budget.get("breakdown") or {})
        budget_rows = [
            ("Personnel", bd.get("personnel", 0)),
            ("Hardware", bd.get("hardware", 0)),
            ("Software", bd.get("software", 0)),
            ("Training", bd.get("training", 0)),
            ("Contingency", bd.get("contingency", 0)),
            ("TOTAL", budget.get("total", 0)),
        ]
        currency = budget.get("currency", "USD")
        add_table(
            doc,
            ["Category", f"Amount ({currency})"],
            [(k, f"{v:,.0f}" if isinstance(v, (int, float)) else v) for k, v in budget_rows],
        )

    # 12. Success Criteria
    doc.add_heading("12. SUCCESS CRITERIA", level=1)
    add_bullet_list(doc, payload.get("success_criteria", []))

    # 13. Approvals
    doc.add_heading("13. APPROVALS", level=1)
    approval_rows = [
        (a.get("role", ""), a.get("name", ""), "", a.get("date", ""))
        for a in payload.get("approvals", [])
    ]
    add_table(doc, ["Role", "Name", "Signature", "Date"], approval_rows)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
