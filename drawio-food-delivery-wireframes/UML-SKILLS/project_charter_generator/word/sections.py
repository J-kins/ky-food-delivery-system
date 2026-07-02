"""Table and list helpers for the Word document builder."""
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_bullet_list(doc, items: list) -> None:
    for item in items or []:
        doc.add_paragraph(str(item), style="List Bullet")


def add_numbered_list(doc, items: list) -> None:
    for item in items or []:
        doc.add_paragraph(str(item), style="List Number")


def add_table(doc, headers: list, rows: list, header_fill: str = "1565C0") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        shading = cell._tc.get_or_add_tcPr()
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        fill = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_fill}"/>')
        shading.append(fill)

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = str(value) if value is not None else ""


def risk_score_color(likelihood: int, impact: int) -> str:
    score = likelihood * impact
    if score >= 15:
        return "FFCDD2"
    if score >= 9:
        return "FFE0B2"
    if score >= 4:
        return "FFF9C4"
    return "C8E6C9"
