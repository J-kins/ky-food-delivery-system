"""Verify generated charter outputs are structurally valid and openable."""
from __future__ import annotations

import logging
import zipfile
from pathlib import Path
from typing import List

log = logging.getLogger(__name__)


def verify_docx(path: str) -> None:
    """Confirm Word document contains native DrawingML shape diagrams."""
    p = Path(path)
    if not p.is_file():
        raise RuntimeError(f"Word document not found: {path}")

    with zipfile.ZipFile(p, "r") as zf:
        names = zf.namelist()
        for req in ("word/document.xml", "[Content_Types].xml"):
            if req not in names:
                raise RuntimeError(f"Invalid DOCX — missing {req}: {path}")
        doc_xml = zf.read("word/document.xml").decode("utf-8", errors="replace")

    shape_groups = doc_xml.count("wordprocessingGroup")
    wsp_shapes = doc_xml.count("wps:wsp")
    if shape_groups < 7:
        raise RuntimeError(
            f"Word document missing editable DrawingML diagrams "
            f"(found {shape_groups} shape groups, expected ≥7): {path}"
        )

    from docx import Document

    doc = Document(str(p))
    if len(doc.paragraphs) < 10:
        raise RuntimeError(f"Word document too sparse ({len(doc.paragraphs)} paragraphs): {path}")
    log.info(
        "Verified DOCX: %s (%d paragraphs, %d shape groups, %d wps:wsp elements)",
        path, len(doc.paragraphs), shape_groups, wsp_shapes,
    )


def verify_vsdx(path: str) -> None:
    """Confirm Visio file is valid ZIP and reloadable by Aspose."""
    p = Path(path)
    if not p.is_file():
        raise RuntimeError(f"Visio file not found: {path}")

    with zipfile.ZipFile(p, "r") as zf:
        names = zf.namelist()
        if not any("visio" in n.lower() for n in names):
            raise RuntimeError(f"Invalid VSDX — no visio/ entries: {path}")

    from diagrams.aspose_renderer import verify_vsdx_readable

    verify_vsdx_readable(str(p))
    log.info("Verified VSDX opens cleanly: %s", path)


def verify_svg_diagrams(svg_dir: str, min_count: int = 1) -> List[str]:
    """Confirm SVG diagram files exist and contain valid XML."""
    d = Path(svg_dir)
    if not d.is_dir():
        raise RuntimeError(f"SVG directory not found: {svg_dir}")
    svgs = sorted(d.glob("*.svg"))
    if len(svgs) < min_count:
        raise RuntimeError(f"Expected at least {min_count} SVG files in {svg_dir}, found {len(svgs)}")
    for svg in svgs:
        text = svg.read_text(encoding="utf-8", errors="replace")
        if "<svg" not in text:
            raise RuntimeError(f"Invalid SVG content: {svg}")
    log.info("Verified %d SVG diagram(s) in %s", len(svgs), svg_dir)
    return [str(s) for s in svgs]


def verify_all_outputs(outputs: dict) -> None:
    """Run all verification checks on build_charter() output dict."""
    if "word" in outputs:
        verify_docx(outputs["word"])
    if "visio" in outputs:
        verify_vsdx(outputs["visio"])
    if "charter_summary" in outputs:
        verify_vsdx(outputs["charter_summary"])
    svg_dir = outputs.get("diagrams_svg_dir")
    if svg_dir:
        verify_svg_diagrams(svg_dir, min_count=1)
    png_dir = Path(svg_dir).parent / "png" if svg_dir else None
    if png_dir and png_dir.is_dir():
        pngs = sorted(png_dir.glob("*.png"))
        if len(pngs) < 7:
            raise RuntimeError(f"Expected ≥7 PNG diagrams in {png_dir}, found {len(pngs)}")
        log.info("Verified %d PNG diagram(s) in %s", len(pngs), png_dir)
